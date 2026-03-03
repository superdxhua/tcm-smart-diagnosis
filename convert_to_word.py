#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档转换脚本
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 读取 Markdown 文件
with open('SUBDOMAIN_ISSUE_REPORT.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = Document()

# 设置文档标题
title = doc.add_heading('子域名"响应数据为空"问题专项报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 解析 Markdown 并添加到 Word
lines = content.split('\n')
in_code_block = False
code_content = []

for line in lines:
    # 处理代码块
    if line.strip().startswith('```'):
        if in_code_block:
            # 结束代码块
            code_para = doc.add_paragraph()
            code_run = code_para.add_run('\n'.join(code_content))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            code_run.font.color.rgb = RGBColor(0, 0, 0)
            code_para.paragraph_format.left_indent = Inches(0.5)
            code_content = []
            in_code_block = False
        else:
            # 开始代码块
            in_code_block = True
            language = line.strip()[3:] or 'text'
            doc.add_paragraph(f'[代码块: {language}]')
        continue

    if in_code_block:
        code_content.append(line)
        continue

    # 处理标题
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)

    # 处理粗体
    elif '**' in line:
        para = doc.add_paragraph()
        parts = re.split(r'\*\*(.+?)\*\*', line)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                para.add_run(part)
            else:
                run = para.add_run(part)
                run.bold = True

    # 处理空行
    elif line.strip() == '':
        doc.add_paragraph()

    # 普通文本
    else:
        # 处理行内代码
        if '`' in line:
            para = doc.add_paragraph()
            parts = re.split(r'`([^`]+)`', line)
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 0, 0)
        else:
            doc.add_paragraph(line)

# 保存 Word 文档
doc.save('SUBDOMAIN_ISSUE_REPORT.docx')
print('Word 文档已生成：SUBDOMAIN_ISSUE_REPORT.docx')
